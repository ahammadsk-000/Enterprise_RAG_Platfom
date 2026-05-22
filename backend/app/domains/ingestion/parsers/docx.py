"""DOCX parser (python-docx). Tables are flattened to tab-separated rows."""

from __future__ import annotations

import io

from app.core.exceptions import ProviderError
from app.domains.ingestion.parsers.base import ParsedDocument, ParsedPage
from app.integrations.ocr.base import OCREngine


class DocxParser:
    mimes = frozenset(
        {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    )

    async def parse(self, content: bytes, *, filename: str, ocr: OCREngine) -> ParsedDocument:
        try:
            import docx  # python-docx  # noqa: PLC0415
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise ProviderError("python-docx is not installed; cannot parse DOCX.") from exc

        document = docx.Document(io.BytesIO(content))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))

        return ParsedDocument.from_pages([ParsedPage(number=1, text="\n".join(parts))])
